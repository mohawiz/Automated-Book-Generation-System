import os
from docx import Document
from db import supabase

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def compile_book():
    # find books ready for final stage
    resp = supabase.table("book_projects") \
        .select("*") \
        .eq("current_stage", "final") \
        .eq("book_output_status", "ready") \
        .execute()

    books = resp.data

    for book in books:
        book_id = book["id"]
        title = book["title"]

        print(f"Compiling book: {title}")

        # fetch chapters in order
        chapters = supabase.table("book_chapters") \
            .select("chapter_number, chapter_title, chapter_text") \
            .eq("book_id", book_id) \
            .order("chapter_number") \
            .execute().data

        if not chapters:
            print("No chapters found, skipping")
            continue

        # create docx
        doc = Document()
        doc.add_heading(title, level=0)

        for ch in chapters:
            doc.add_heading(
                f"Chapter {ch['chapter_number']}: {ch['chapter_title']}",
                level=1
            )
            doc.add_paragraph(ch["chapter_text"])

        file_name = f"{title.replace(' ', '_')}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)
        doc.save(file_path)

        # upload to Supabase Storage
        with open(file_path, "rb") as f:
            supabase.storage.from_("books").upload(
                path=file_name,
                file=f,
                file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            )

        # update DB
        supabase.table("book_projects").update({
            "book_output_status": "compiled",
            "current_stage": "completed"
        }).eq("id", book_id).execute()

        print("Final book compiled and uploaded ✅")